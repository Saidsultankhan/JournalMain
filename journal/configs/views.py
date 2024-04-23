from rest_framework import viewsets, views
from django.http import HttpResponse
from django.core.exceptions import PermissionDenied
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
    Frame,
)
from configs.models import (
    Grade,
    DairyOfClass,
    User,
    SubjectTeacher,
    Conversation,
    Message,
    Homework,
    Role,
)

# from users.models import (
#     Role,
#     User
# )
from configs.serializers import (
    GradeDetailSerializer,
    GradeListSerializer,
    GradeSerializer,
    DairyListSerializer,
    DairyDetailSerializer,
    DairyDefaultSerializer,
    DairyCreateSerializer,
    TeacherCreateSerializer,
    TeacherDetailSerializer,
    TeacherListSerializer,
    TeacherDefaultSerializer,
    PupilDetailSerializer,
    PupilListSerializer,
    PupilDefaultSerializers,
    SubjectTeacherListSerializer,
    SubjectTeacherDetailSerializer,
    SubjectTeacherDefaultSerializer,
    UserSerializer,
    ConversationSerializer,
    # UserCreateUpdateSerializer,
    MessageSerializer,
    MessageCreateUpdateSerializer,
    HomeworkDefaultSerializer,
    HomeworkCreateSerializer,
    HomeworkStudentUpdateSerializer,
    HomeworkTeacherUpdateSerializer,
)
from rest_framework.permissions import IsAdminUser, IsAuthenticated, AllowAny
from configs.permissions import (
    GradeDetailPermisson,
    GradeListPermission,
    DairyListTeacher,
    DairyTeacher,
    PupilDetailPermission,
    SubjectTeacherPermission,
)
from rest_framework.views import APIView
import os
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert
from openpyxl import Workbook
from openpyxl.styles import Font, Side, Border, Alignment
from openpyxl.utils import get_column_letter


class GradeViewSet(viewsets.ModelViewSet):
    queryset = Grade.objects.all()

    def get_serializer_class(self):
        serializers = {
            "retrieve": GradeDetailSerializer,
            "create": GradeSerializer,
            "list": GradeListSerializer,
            "delete": GradeSerializer,
            "partial_update": GradeSerializer,
        }
        return serializers.get(self.action, GradeSerializer)

    def get_permissions(self):
        permission_classes = []
        if self.action == "retrieve":
            permission_classes = [IsAuthenticated and GradeDetailPermisson]
        elif self.action == "list":
            permission_classes = [GradeListPermission]
        elif self.action in ["create", "partial_update", "destroy"]:
            permission_classes = [IsAdminUser]
        return [permission_class() for permission_class in permission_classes]


class DairyViewSet(viewsets.ModelViewSet):
    queryset = DairyOfClass.objects.all()

    def get_serializer_class(self):
        serializers = {
            "retrieve": DairyDetailSerializer,
            "create": DairyCreateSerializer,
            "list": DairyListSerializer,
            "delete": DairyDefaultSerializer,
            "partial_update": DairyDefaultSerializer,
            "update": DairyDefaultSerializer,
        }
        return serializers.get(self.action, DairyDefaultSerializer)

    def get_permissions(self):
        permission_classes = []
        if self.action in ["retrieve", "create", "partial_update", "destroy", "update"]:
            permission_classes = [DairyTeacher]
        elif self.action in ["list"]:
            permission_classes = [IsAuthenticated, DairyListTeacher]

        return [permission_class() for permission_class in permission_classes]


class TeacherViewSet(viewsets.ModelViewSet):
    queryset = User.objects.filter(role__user_type=Role.TEACHER)

    def get_serializer_class(self):
        serializers = {
            "list": TeacherListSerializer,
            "retrieve": TeacherDetailSerializer,
            "create": TeacherCreateSerializer,
            "partial_update": TeacherDefaultSerializer,
            "delete": TeacherDefaultSerializer,
        }
        return serializers.get(self.action, TeacherDefaultSerializer)

    def get_permissions(self):
        permission_classes = []
        if self.action in ["retrieve", "create", "partial_update", "destroy", "update"]:
            permission_classes = [IsAdminUser]
        elif self.action == "list":
            permission_classes = [AllowAny]

        return [permission_class() for permission_class in permission_classes]


class PupilsViewSet(viewsets.ModelViewSet):
    queryset = User.objects.filter(role__user_type=Role.STUDENT)

    def get_serializer_class(self):
        serializers = {
            "retrieve": PupilDetailSerializer,
            "list": PupilListSerializer,
            "partial_update": PupilDefaultSerializers,
            "create": PupilDefaultSerializers,
            "delete": PupilDefaultSerializers,
        }
        return serializers.get(self.action, PupilDetailSerializer)

    def get_permissions(self):
        permission_classes = []
        if self.action == "retrieve":
            permission_classes = [PupilDetailPermission]
        elif self.action == "list":
            permission_classes = [IsAuthenticated, IsAdminUser or PupilDetailPermission]
        elif self.action in ["create", "partial_update", "destroy"]:
            permission_classes = [IsAdminUser]

        return [permission_class() for permission_class in permission_classes]


class PDFGeneratorView(views.APIView):

    def get(self, request):
        filename = "marks.pdf"
        buffer = BytesIO()

        pdf = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        header_style = styles["Heading1"]

        database = [["Name", "Grade", "Subject", "Quarter", "mark"]]
        pupils_data = DairyOfClass.objects.filter(pupil=request.user).prefetch_related(
            "pupil", "subject", "grade"
        )
        user = request.user

        for i in pupils_data:
            database.append(
                [
                    f"{user}",
                    f"{str(i.grade)}",
                    f"{str(i.subject)}",
                    f"{i.quarter}",
                    f"{i.mark}",
                ]
            )

        # for dairy_data in obj.dairy_pupil.values():
        #     print(dairy_data['mark'])
        # print(obj.dairy_pupil.values())
        header_text = f"Butun yil davomida olingan {user}-ning baholari."
        header_frame = Frame(pdf.width - 200, pdf.height - 50, 200, 50, showBoundary=0)

        table = Table(database)
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )

        table.setStyle(style)

        elements = []

        header_content = Paragraph(header_text, header_style)
        elements.append(header_content)
        elements.append(table)
        header_frame.addFromList([header_content], pdf)
        pdf.build(elements)

        pdf_data = buffer.getvalue()
        buffer.close()

        response = HttpResponse(content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        response.write(pdf_data)

        return response


# libreoffice_path = "/Applications/libreoffice/Contents/MacOS"
# os.environ['LIBRE_OFFICE_PATH'] = libreoffice_path


class WordToPDFView(APIView):
    def get(self, request):
        database = []
        pupils_data = DairyOfClass.objects.filter(pupil=request.user).prefetch_related(
            "pupil", "subject", "grade"
        )
        user = request.user

        for pupil in pupils_data:
            database.append(
                [
                    f"{user}",
                    f"{str(pupil.grade)}",
                    f"{str(pupil.subject)}",
                    f"{pupil.quarter}",
                    f"{pupil.mark}",
                ]
            )
        data = {
            "headers": ["Name", "Grade", "Subject", "Quarter", "Mark"],
            "rows": database,
        }

        document = Document()
        document.add_paragraph(
            f"Butun yil davomida olingan {user}-ning baholari.", style="Heading1"
        )

        table = document.add_table(rows=1, cols=len(data["headers"]))
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(data["headers"]):
            hdr_cells[idx].text = header

        for row in data["rows"]:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = value

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        with open("temporary.docx", "wb") as f:
            f.write(buffer.getvalue())

        output_file = f"{user}-ning baho tabeli.pdf"

        try:
            convert("temporary.docx", f"{output_file}")
        except Exception as e:
            print(f"Ошибка при конвертации: {e}")

        with open(output_file, "rb") as pdf_file:
            response = HttpResponse(pdf_file.read(), content_type="application/pdf")
            response["Content-Disposition"] = f"inline; filename={output_file}"
            return response


class ExcelGenerateViewSet(views.APIView):  # retrieveapiview

    def get(self, request):
        wb = Workbook()
        wb.remove(wb.active)

        pupils_sheet = wb.create_sheet("Ученики", 0)
        pupils_sheet.append(
            [
                "№",
                "Имя_ru",
                "Имя_en",
                "Имя_uz",
                "Класс",
                "Классный руководитель",
                "Родители",
            ]
        )
        number = 1

        for obj in User.objects.filter(role__user_type=Role.STUDENT).prefetch_related(
            "dairy_pupil", "grade", "parent"
        ):
            pupils_sheet.append(
                [
                    number,
                    obj.name_ru,
                    obj.name_en,
                    obj.name_uz,
                    str(obj.grade),
                    obj.grade.teacher.name_uz,
                    # obj.parent.name_uz
                ]
            )
            number += 1
        max_lengths = {}
        first_iteration = True

        for row in pupils_sheet.iter_rows(
            min_row=1,
            max_row=pupils_sheet.max_row,
            min_col=1,
            max_col=pupils_sheet.max_column,
        ):
            for cell in row:
                if first_iteration:
                    cell.style = "Good"
                    cell.font = cell.font.copy(bold=True, size=14, color="000000")
                pupils_sheet[cell.coordinate].alignment = Alignment(horizontal="center")
                row[0].font = Font(bold=True, size=14)

                if cell.value:
                    length = len(str(cell.value))
                    if (
                        cell.column_letter not in max_lengths
                        or length > max_lengths[cell.column_letter]
                    ):
                        max_lengths[cell.column_letter] = length
                cell.border = Border(
                    left=Side(border_style="thin", color="000000"),
                    right=Side(border_style="thin", color="000000"),
                    top=Side(border_style="thin", color="000000"),
                    bottom=Side(border_style="thin", color="000000"),
                )
            first_iteration = False

        for col, length in max_lengths.items():
            pupils_sheet.column_dimensions[col].width = length + 10

        file_name = "pupils_detail.xlsx"

        wb.save(file_name)

        with open(file_name, "rb") as excel_file:
            response = HttpResponse(
                excel_file.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = f"attachment; filename={file_name}"
            return response


class SubjectTeacherViewSet(viewsets.ModelViewSet):

    def get_serializer_class(self):
        serializers = {
            "list": SubjectTeacherListSerializer,
            "retrieve": SubjectTeacherDetailSerializer,
            "create": SubjectTeacherDefaultSerializer,
            "partial_update": SubjectTeacherDefaultSerializer,
            "delete": SubjectTeacherDefaultSerializer,
        }
        return serializers.get(self.action, SubjectTeacherDetailSerializer)

    def get_permissions(self):
        permission_classes = []
        if self.action in ["retrieve", "list"]:
            permission_classes = [IsAdminUser | SubjectTeacherPermission]
        elif self.action in ["create", "partial_update", "destroy"]:
            permission_classes = [IsAdminUser]

        return [permission_class() for permission_class in permission_classes]

    def get_queryset(self):
        is_mentor = self.request.query_params.get(
            "mentor_id", False
        )  # used it just to test the filters
        teacher = self.request.query_params.get("teacher_id")

        user = self.request.user
        if not user.is_authenticated:
            raise PermissionDenied()
        elif user.is_staff:
            return SubjectTeacher.objects.all().order_by("id")
        elif user.role.user_type == Role.TEACHER:
            return SubjectTeacher.objects.filter(teacher=teacher)

    # def get_serializer_class(self):
    #     serializers = {
    #         'create': UserCreateUpdateSerializer,
    #         'partial_update': UserCreateUpdateSerializer,
    #     }

    #     return serializers.get(self.action, UserSerializer)


class UserModelViewset(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer


class ConversationViewSet(viewsets.ModelViewSet):
    queryset = Conversation.objects.all()
    serializer_class = ConversationSerializer

    def get_queryset(self):
        user = self.request.user

        if not user.is_staff:
            user_conversations = Conversation.objects.filter(participant=user)
            queryset = Message.objects.filter(conversation__in=user_conversations)

            return queryset

        if user.is_staff:
            queryset = Message.objects.all().order_by("conversation")

        return ValueError("Что-то пошло не так")


class MessageViewSet(viewsets.ModelViewSet):
    serializer_class = MessageSerializer

    def get_serializer_class(self):
        serializers = {
            "create": MessageCreateUpdateSerializer,
        }

        return serializers.get(self.action, MessageCreateUpdateSerializer)

    def get_queryset(self):

        if not self.request.user.id:
            raise PermissionDenied("tasssssss")

        admin = self.request.user.is_staff
        user = self.request.user

        if admin:
            queryset = Message.objects.all().order_by("conversation")
            return queryset
        if user:
            queryset = Message.objects.filter(sender=user).order_by("conversation")
            return queryset

        return PermissionDenied("no-no-no")


class HomeworkViewSet(viewsets.ModelViewSet):

    serializer_class = HomeworkDefaultSerializer
    queryset = Homework.objects.all()

    def get_serializer_class(self):
        user_role = self.request.user.role.user_type

        if user_role == Role.TEACHER:
            serializers = {
                "create": HomeworkCreateSerializer,
                "partial_update": HomeworkTeacherUpdateSerializer,
                "list": HomeworkDefaultSerializer,
            }
            
            return serializers.get(self.action, HomeworkDefaultSerializer)

        elif user_role == Role.STUDENT:
            serializers = {
                'partial_update': HomeworkStudentUpdateSerializer,
            }
            print(self.action)
            return serializers.get(self.action, HomeworkDefaultSerializer)
        
        return PermissionDenied("no-no-no")
